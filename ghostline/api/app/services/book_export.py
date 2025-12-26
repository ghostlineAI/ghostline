"""
Book Export Service - Converts generated books to multiple output formats.

Supported formats:
- PDF: Professional print-ready documents
- EPUB: E-reader compatible format
- DOCX: Microsoft Word format
- TXT: Plain text
- HTML: Web-ready format
- Markdown: Source format for further editing
"""

import io
import os
import re
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Optional, Union
from dataclasses import dataclass

from pydantic import BaseModel


class ExportFormat(str, Enum):
    """Supported export formats."""
    PDF = "pdf"
    EPUB = "epub"
    DOCX = "docx"
    TXT = "txt"
    HTML = "html"
    MARKDOWN = "md"


@dataclass
class BookMetadata:
    """Metadata for the exported book."""
    title: str
    author: str = "Anonymous"
    description: str = ""
    language: str = "en"
    publisher: str = "GhostLine"
    created_date: str = ""
    keywords: list[str] = None
    
    def __post_init__(self):
        if not self.created_date:
            self.created_date = datetime.now().strftime("%Y-%m-%d")
        if self.keywords is None:
            self.keywords = []


@dataclass
class CitationMetadata:
    """Metadata for a citation in the text."""
    filename: str  # Source file name
    quote: str  # The verbatim quoted text
    quote_start: int = None  # Start position in clean text
    quote_end: int = None  # End position in clean text
    source_material_id: str = None  # DB reference for future linking
    verified: bool = False  # Whether quote was verified against source


@dataclass
class Chapter:
    """A book chapter."""
    number: int
    title: str
    content: str  # Clean content without citation markers
    word_count: int = 0
    citations: list[CitationMetadata] = None  # Citation metadata for references section
    
    def __post_init__(self):
        if not self.word_count:
            self.word_count = len(self.content.split())
        if self.citations is None:
            self.citations = []


class BookExportService:
    """
    Service for exporting books to various formats.
    
    Usage:
        service = BookExportService()
        
        # Export to PDF
        pdf_bytes = service.export_pdf(chapters, metadata)
        
        # Export to all formats
        exports = service.export_all(chapters, metadata, output_dir="/path/to/output")
    """
    
    def __init__(self, output_dir: Optional[str] = None):
        """
        Initialize the export service.
        
        Args:
            output_dir: Default output directory for exports
        """
        self.output_dir = Path(output_dir) if output_dir else Path.cwd() / "exports"
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def export(
        self,
        chapters: list[Chapter],
        metadata: BookMetadata,
        format: ExportFormat,
        output_path: Optional[str] = None,
    ) -> Union[bytes, str]:
        """
        Export book to specified format.
        
        Args:
            chapters: List of Chapter objects
            metadata: Book metadata
            format: Target export format
            output_path: Optional path to save file
            
        Returns:
            bytes for binary formats (PDF, EPUB, DOCX), str for text formats
        """
        exporters = {
            ExportFormat.PDF: self._export_pdf,
            ExportFormat.EPUB: self._export_epub,
            ExportFormat.DOCX: self._export_docx,
            ExportFormat.TXT: self._export_txt,
            ExportFormat.HTML: self._export_html,
            ExportFormat.MARKDOWN: self._export_markdown,
        }
        
        exporter = exporters.get(format)
        if not exporter:
            raise ValueError(f"Unsupported format: {format}")
        
        result = exporter(chapters, metadata)
        
        if output_path:
            self._save_to_file(result, output_path, format)
        
        return result
    
    def export_all(
        self,
        chapters: list[Chapter],
        metadata: BookMetadata,
        output_dir: Optional[str] = None,
        base_filename: Optional[str] = None,
    ) -> dict[ExportFormat, str]:
        """
        Export book to all supported formats.
        
        Args:
            chapters: List of Chapter objects
            metadata: Book metadata
            output_dir: Directory to save exports
            base_filename: Base name for output files (without extension)
            
        Returns:
            Dict mapping format to output file path
        """
        output_dir = Path(output_dir) if output_dir else self.output_dir
        output_dir.mkdir(parents=True, exist_ok=True)
        
        if not base_filename:
            # Create safe filename from title
            base_filename = self._sanitize_filename(metadata.title)
        
        results = {}
        
        for format in ExportFormat:
            try:
                output_path = output_dir / f"{base_filename}.{format.value}"
                self.export(chapters, metadata, format, str(output_path))
                results[format] = str(output_path)
                print(f"✅ Exported {format.value.upper()}: {output_path}")
            except Exception as e:
                print(f"❌ Failed to export {format.value.upper()}: {e}")
                results[format] = f"ERROR: {e}"
        
        return results
    
    def _sanitize_filename(self, name: str) -> str:
        """Convert title to safe filename."""
        # Remove special characters, replace spaces with underscores
        safe = re.sub(r'[^\w\s-]', '', name)
        safe = re.sub(r'[-\s]+', '_', safe)
        return safe.lower()[:50]  # Limit length
    
    def _save_to_file(self, content: Union[bytes, str], path: str, format: ExportFormat):
        """Save content to file."""
        mode = 'wb' if isinstance(content, bytes) else 'w'
        encoding = None if isinstance(content, bytes) else 'utf-8'
        
        with open(path, mode, encoding=encoding) as f:
            f.write(content)

    # =========================================================================
    # Citation helpers
    # =========================================================================

    _CITATION_PATTERN = re.compile(r"\[citation:\s*([^\]]+)\]", re.IGNORECASE)

    def _extract_citations_as_footnotes(
        self,
        content: str,
        marker_style: str = "none",  # retained for backwards compat; we no longer emit inline markers
    ) -> tuple[str, list[str]]:
        """
        Remove inline [citation: ...] markers from the text and return the cleaned text
        plus a list of note bodies (in first-appearance order).

        IMPORTANT:
        - We do NOT emit inline numeric markers (e.g. [1]) anymore. The product requirement
          is to keep citations out of the prose and render them in a Notes/References section.
        """
        if not content:
            return content, []

        note_numbers: dict[str, int] = {}
        notes: list[str] = []

        def _repl(match: re.Match) -> str:
            note_text = (match.group(1) or "").strip()
            if not note_text:
                return ""
            if note_text not in note_numbers:
                note_numbers[note_text] = len(notes) + 1
                notes.append(note_text)
            # Do not insert any inline footnote markers in the prose.
            return ""

        processed = self._CITATION_PATTERN.sub(_repl, content)
        processed = re.sub(r"[ \t]{2,}", " ", processed)
        return processed, notes

    def _render_html_notes(self, notes: list[str]) -> str:
        """Render an end-of-chapter notes block for HTML/EPUB."""
        if not notes:
            return ""
        # Basic escaping (we already keep these as plain text strings)
        def esc(s: str) -> str:
            return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        items = "\n".join(f"<li>[{i}] {esc(note)}</li>" for i, note in enumerate(notes, 1))
        return f"""
        <section class="footnotes">
            <h2>Notes</h2>
            <ol>
                {items}
            </ol>
        </section>
        """
    
    # =========================================================================
    # PDF Export
    # =========================================================================
    
    def _export_pdf(self, chapters: list[Chapter], metadata: BookMetadata) -> bytes:
        """Export to PDF format using reportlab."""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, PageBreak,
                Table, TableStyle
            )
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
            from reportlab.lib import colors
        except ImportError:
            # Fallback: create a simple text-based PDF using basic reportlab
            return self._export_pdf_simple(chapters, metadata)
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )
        
        # Styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'BookTitle',
            parent=styles['Heading1'],
            fontSize=28,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#1a1a2e'),
        )
        
        author_style = ParagraphStyle(
            'Author',
            parent=styles['Normal'],
            fontSize=16,
            spaceAfter=50,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#4a4a6a'),
        )
        
        chapter_title_style = ParagraphStyle(
            'ChapterTitle',
            parent=styles['Heading1'],
            fontSize=20,
            spaceBefore=20,
            spaceAfter=20,
            textColor=colors.HexColor('#1a1a2e'),
        )
        
        body_style = ParagraphStyle(
            'BookBody',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            spaceAfter=12,
            alignment=TA_JUSTIFY,
        )
        
        heading2_style = ParagraphStyle(
            'BookHeading2',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=16,
            spaceAfter=10,
            textColor=colors.HexColor('#2a2a4e'),
        )
        
        # Build document
        story = []
        
        # Title page
        story.append(Spacer(1, 2 * inch))
        story.append(Paragraph(metadata.title, title_style))
        story.append(Paragraph(f"by {metadata.author}", author_style))
        if metadata.description:
            desc_style = ParagraphStyle(
                'Description',
                parent=styles['Normal'],
                fontSize=12,
                alignment=TA_CENTER,
                textColor=colors.HexColor('#666666'),
            )
            story.append(Paragraph(metadata.description, desc_style))
        story.append(Spacer(1, 2 * inch))
        story.append(Paragraph(f"Published by {metadata.publisher}", author_style))
        story.append(Paragraph(metadata.created_date, author_style))
        story.append(PageBreak())
        
        # Table of Contents
        toc_title = ParagraphStyle(
            'TOCTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
            alignment=TA_CENTER,
        )
        story.append(Paragraph("Table of Contents", toc_title))
        story.append(Spacer(1, 20))
        
        for chapter in chapters:
            toc_entry = ParagraphStyle(
                'TOCEntry',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=8,
            )
            story.append(Paragraph(
                f"Chapter {chapter.number}: {chapter.title}",
                toc_entry
            ))
        
        story.append(PageBreak())
        
        # Chapters
        for chapter in chapters:
            # Chapter title
            story.append(Paragraph(
                f"Chapter {chapter.number}: {chapter.title}",
                chapter_title_style
            ))
            story.append(Spacer(1, 20))
            
            # Clean content (remove any residual inline citation markers)
            chapter_body, marker_notes = self._extract_citations_as_footnotes(
                chapter.content, marker_style="none"
            )

            # Prefer structured citation metadata if present; fall back to marker notes.
            notes = []
            if chapter.citations:
                for cit in chapter.citations:
                    filename = (cit.filename or "Unknown").strip()
                    quote = (cit.quote or "").strip()
                    if quote:
                        notes.append(f"{filename} - \"{quote}\"")
                    else:
                        notes.append(filename)
            else:
                notes = marker_notes

            # Process content - convert markdown to paragraphs
            paragraphs = self._markdown_to_paragraphs(
                chapter_body, body_style, heading2_style
            )
            story.extend(paragraphs)

            # Notes (end-of-chapter; no inline markers)
            if notes:
                story.append(Spacer(1, 12))
                story.append(Paragraph("Notes", heading2_style))
                note_style = ParagraphStyle(
                    "Footnote",
                    parent=body_style,
                    fontSize=9,
                    leading=12,
                    textColor=colors.HexColor("#444444"),
                    spaceAfter=6,
                )
                for i, note in enumerate(notes, 1):
                    safe_note = note.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    story.append(Paragraph(f"[{i}] {safe_note}", note_style))
            
            story.append(PageBreak())
        
        doc.build(story)
        buffer.seek(0)
        return buffer.read()
    
    def _export_pdf_simple(self, chapters: list[Chapter], metadata: BookMetadata) -> bytes:
        """Simple PDF export fallback."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
        except ImportError:
            raise ImportError("reportlab is required for PDF export. Install with: pip install reportlab")
        
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        
        # Title page
        c.setFont("Helvetica-Bold", 24)
        c.drawCentredString(width/2, height - 200, metadata.title)
        c.setFont("Helvetica", 14)
        c.drawCentredString(width/2, height - 250, f"by {metadata.author}")
        c.showPage()
        
        # Chapters
        for chapter in chapters:
            y = height - 72
            c.setFont("Helvetica-Bold", 16)
            c.drawString(72, y, f"Chapter {chapter.number}: {chapter.title}")
            y -= 30
            
            c.setFont("Helvetica", 11)
            # Simple text wrapping
            lines = self._wrap_text(chapter.content, 80)
            for line in lines:
                if y < 72:
                    c.showPage()
                    y = height - 72
                c.drawString(72, y, line)
                y -= 14
            
            c.showPage()
        
        c.save()
        buffer.seek(0)
        return buffer.read()
    
    def _markdown_to_paragraphs(self, content: str, body_style, heading_style) -> list:
        """Convert markdown content to reportlab paragraphs."""
        from reportlab.platypus import Paragraph, Spacer
        
        paragraphs = []
        lines = content.split('\n')
        current_para = []
        
        for line in lines:
            line = line.strip()
            
            if not line:
                if current_para:
                    text = ' '.join(current_para)
                    # Escape special characters for reportlab
                    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    # Handle bold/italic markdown
                    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
                    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
                    paragraphs.append(Paragraph(text, body_style))
                    current_para = []
                continue
            
            # Heading detection
            if line.startswith('## '):
                if current_para:
                    text = ' '.join(current_para)
                    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    paragraphs.append(Paragraph(text, body_style))
                    current_para = []
                heading_text = line[3:].strip()
                paragraphs.append(Spacer(1, 10))
                paragraphs.append(Paragraph(heading_text, heading_style))
                continue
            
            # Bold lines (like **F - Find Your Breath**)
            if line.startswith('**') and line.endswith('**'):
                if current_para:
                    text = ' '.join(current_para)
                    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    paragraphs.append(Paragraph(text, body_style))
                    current_para = []
                bold_text = line[2:-2]
                bold_text = bold_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                paragraphs.append(Paragraph(f"<b>{bold_text}</b>", body_style))
                continue
            
            current_para.append(line)
        
        # Handle remaining paragraph
        if current_para:
            text = ' '.join(current_para)
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
            text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
            paragraphs.append(Paragraph(text, body_style))
        
        return paragraphs
    
    def _wrap_text(self, text: str, width: int) -> list[str]:
        """Simple text wrapping."""
        words = text.replace('\n', ' ').split()
        lines = []
        current_line = []
        current_length = 0
        
        for word in words:
            if current_length + len(word) + 1 <= width:
                current_line.append(word)
                current_length += len(word) + 1
            else:
                if current_line:
                    lines.append(' '.join(current_line))
                current_line = [word]
                current_length = len(word)
        
        if current_line:
            lines.append(' '.join(current_line))
        
        return lines
    
    # =========================================================================
    # EPUB Export
    # =========================================================================
    
    def _export_epub(self, chapters: list[Chapter], metadata: BookMetadata) -> bytes:
        """Export to EPUB format."""
        try:
            from ebooklib import epub
        except ImportError:
            raise ImportError("ebooklib is required for EPUB export. Install with: pip install ebooklib")
        
        book = epub.EpubBook()
        
        # Set metadata
        book.set_identifier(f'ghostline-{metadata.title.lower().replace(" ", "-")}')
        book.set_title(metadata.title)
        book.set_language(metadata.language)
        book.add_author(metadata.author)
        
        # Add metadata
        book.add_metadata('DC', 'publisher', metadata.publisher)
        book.add_metadata('DC', 'description', metadata.description)
        book.add_metadata('DC', 'date', metadata.created_date)
        
        # CSS for styling
        style = '''
        @namespace epub "http://www.idpf.org/2007/ops";
        body {
            font-family: Georgia, serif;
            line-height: 1.6;
            margin: 5%;
        }
        h1 {
            text-align: center;
            color: #1a1a2e;
            margin-bottom: 2em;
        }
        h2 {
            color: #2a2a4e;
            margin-top: 1.5em;
        }
        p {
            text-indent: 1em;
            margin: 0.5em 0;
        }
        .chapter-title {
            font-size: 1.5em;
            margin-bottom: 1em;
        }
        '''
        
        css = epub.EpubItem(
            uid="style",
            file_name="style/style.css",
            media_type="text/css",
            content=style
        )
        book.add_item(css)
        
        # Create chapters
        epub_chapters = []
        
        for chapter in chapters:
            # Convert markdown to HTML
            chapter_body, marker_notes = self._extract_citations_as_footnotes(
                chapter.content, marker_style="none"
            )

            notes = []
            if chapter.citations:
                for cit in chapter.citations:
                    filename = (cit.filename or "Unknown").strip()
                    quote = (cit.quote or "").strip()
                    notes.append(f"{filename} - \"{quote}\"" if quote else filename)
            else:
                notes = marker_notes

            html_content = self._markdown_to_html(chapter_body) + self._render_html_notes(notes)
            
            c = epub.EpubHtml(
                title=f"Chapter {chapter.number}: {chapter.title}",
                file_name=f'chapter_{chapter.number}.xhtml',
                lang=metadata.language
            )
            c.content = f'''
            <html>
            <head>
                <link rel="stylesheet" href="style/style.css" type="text/css"/>
            </head>
            <body>
                <h1 class="chapter-title">Chapter {chapter.number}: {chapter.title}</h1>
                {html_content}
            </body>
            </html>
            '''
            c.add_item(css)
            book.add_item(c)
            epub_chapters.append(c)
        
        # Create table of contents
        book.toc = [(epub.Section('Chapters'), epub_chapters)]
        
        # Add navigation files
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Create spine
        book.spine = ['nav'] + epub_chapters
        
        # Write to bytes
        buffer = io.BytesIO()
        epub.write_epub(buffer, book)
        buffer.seek(0)
        return buffer.read()
    
    def _markdown_to_html(self, content: str) -> str:
        """Convert markdown to HTML."""
        html_lines = []
        lines = content.split('\n')
        in_paragraph = False
        current_para = []
        
        for line in lines:
            stripped = line.strip()
            
            if not stripped:
                if current_para:
                    para_text = ' '.join(current_para)
                    para_text = self._convert_inline_markdown(para_text)
                    html_lines.append(f'<p>{para_text}</p>')
                    current_para = []
                continue
            
            # Headings
            if stripped.startswith('## '):
                if current_para:
                    para_text = ' '.join(current_para)
                    para_text = self._convert_inline_markdown(para_text)
                    html_lines.append(f'<p>{para_text}</p>')
                    current_para = []
                heading = stripped[3:]
                html_lines.append(f'<h2>{heading}</h2>')
                continue
            
            # Bold standalone lines
            if stripped.startswith('**') and stripped.endswith('**') and stripped.count('**') == 2:
                if current_para:
                    para_text = ' '.join(current_para)
                    para_text = self._convert_inline_markdown(para_text)
                    html_lines.append(f'<p>{para_text}</p>')
                    current_para = []
                bold_text = stripped[2:-2]
                html_lines.append(f'<p><strong>{bold_text}</strong></p>')
                continue
            
            current_para.append(stripped)
        
        if current_para:
            para_text = ' '.join(current_para)
            para_text = self._convert_inline_markdown(para_text)
            html_lines.append(f'<p>{para_text}</p>')
        
        return '\n'.join(html_lines)
    
    def _convert_inline_markdown(self, text: str) -> str:
        """Convert inline markdown (bold, italic) to HTML."""
        # Bold
        text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
        # Italic
        text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)
        # Citations - style them nicely
        text = re.sub(r'\[citation: ([^\]]+)\]', r'<em class="citation">[\1]</em>', text)
        return text
    
    # =========================================================================
    # DOCX Export
    # =========================================================================
    
    def _export_docx(self, chapters: list[Chapter], metadata: BookMetadata) -> bytes:
        """Export to DOCX (Microsoft Word) format."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")
        
        doc = Document()
        
        # Title
        title = doc.add_heading(metadata.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Author
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = author_para.add_run(f"by {metadata.author}")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Description
        if metadata.description:
            desc_para = doc.add_paragraph()
            desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = desc_para.add_run(metadata.description)
            run.font.size = Pt(11)
            run.font.italic = True
        
        # Publisher info
        doc.add_paragraph()
        pub_para = doc.add_paragraph()
        pub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pub_para.add_run(f"{metadata.publisher} • {metadata.created_date}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_page_break()
        
        # Table of Contents
        doc.add_heading("Table of Contents", 1)
        for chapter in chapters:
            toc_para = doc.add_paragraph()
            run = toc_para.add_run(f"Chapter {chapter.number}: {chapter.title}")
            run.font.size = Pt(11)
        
        doc.add_page_break()
        
        # Chapters
        for chapter in chapters:
            doc.add_heading(f"Chapter {chapter.number}: {chapter.title}", 1)
            
            # Process content
            chapter_body, marker_notes = self._extract_citations_as_footnotes(
                chapter.content, marker_style="none"
            )
            notes = []
            if chapter.citations:
                for cit in chapter.citations:
                    filename = (cit.filename or "Unknown").strip()
                    quote = (cit.quote or "").strip()
                    notes.append(f"{filename} - \"{quote}\"" if quote else filename)
            else:
                notes = marker_notes
            self._add_markdown_to_docx(doc, chapter_body)

            # Notes (end-of-chapter; no inline markers)
            if notes:
                doc.add_heading("Notes", 2)
                for i, note in enumerate(notes, 1):
                    doc.add_paragraph(f"[{i}] {note}")
            
            doc.add_page_break()
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def _add_markdown_to_docx(self, doc, content: str):
        """Add markdown content to docx document."""
        from docx.shared import Pt
        
        lines = content.split('\n')
        current_para = []
        
        for line in lines:
            stripped = line.strip()
            
            if not stripped:
                if current_para:
                    para_text = ' '.join(current_para)
                    self._add_formatted_paragraph(doc, para_text)
                    current_para = []
                continue
            
            # Headings
            if stripped.startswith('## '):
                if current_para:
                    para_text = ' '.join(current_para)
                    self._add_formatted_paragraph(doc, para_text)
                    current_para = []
                heading = stripped[3:]
                doc.add_heading(heading, 2)
                continue
            
            # Bold standalone lines
            if stripped.startswith('**') and stripped.endswith('**') and stripped.count('**') == 2:
                if current_para:
                    para_text = ' '.join(current_para)
                    self._add_formatted_paragraph(doc, para_text)
                    current_para = []
                bold_text = stripped[2:-2]
                para = doc.add_paragraph()
                run = para.add_run(bold_text)
                run.bold = True
                continue
            
            current_para.append(stripped)
        
        if current_para:
            para_text = ' '.join(current_para)
            self._add_formatted_paragraph(doc, para_text)
    
    def _add_formatted_paragraph(self, doc, text: str):
        """Add a paragraph with inline formatting."""
        from docx.shared import RGBColor
        
        para = doc.add_paragraph()
        
        # Parse inline formatting
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|\[citation:[^\]]+\])', text)
        
        for part in parts:
            if not part:
                continue
            
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = para.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('[citation:'):
                run = para.add_run(part)
                run.italic = True
                run.font.color.rgb = RGBColor(100, 100, 150)
            else:
                para.add_run(part)
    
    # =========================================================================
    # TXT Export
    # =========================================================================
    
    def _export_txt(self, chapters: list[Chapter], metadata: BookMetadata) -> str:
        """Export to plain text format with references section."""
        lines = []
        
        # Title
        lines.append("=" * 60)
        lines.append(metadata.title.center(60))
        lines.append(f"by {metadata.author}".center(60))
        lines.append("=" * 60)
        lines.append("")
        
        if metadata.description:
            lines.append(metadata.description)
            lines.append("")
        
        lines.append(f"Published by {metadata.publisher}")
        lines.append(f"Date: {metadata.created_date}")
        lines.append("")
        lines.append("-" * 60)
        lines.append("")
        
        # Table of Contents
        lines.append("TABLE OF CONTENTS")
        lines.append("-" * 20)
        for chapter in chapters:
            lines.append(f"  Chapter {chapter.number}: {chapter.title}")
        lines.append("")
        lines.append("=" * 60)
        lines.append("")
        
        # Chapters - use clean content (no inline citation markers)
        all_references = []  # Collect all references for end-of-book section
        
        for chapter in chapters:
            lines.append("-" * 60)
            lines.append(f"CHAPTER {chapter.number}: {chapter.title.upper()}")
            lines.append("-" * 60)
            lines.append("")
            
            # The content should already be clean (no citation markers)
            # Just clean up markdown formatting
            content = chapter.content
            content = re.sub(r'\*\*(.+?)\*\*', r'\1', content)  # Remove bold
            content = re.sub(r'\*(.+?)\*', r'\1', content)  # Remove italic
            content = re.sub(r'^## ', '', content, flags=re.MULTILINE)  # Remove heading markers
            # Remove any remaining citation markers (in case content wasn't pre-cleaned)
            content = re.sub(r'\[citation:\s*[^\]]+\]', '', content, flags=re.IGNORECASE)
            content = re.sub(r'  +', ' ', content)  # Clean double spaces
            
            lines.append(content.strip())
            
            # Collect citations for this chapter into references
            if chapter.citations:
                for cit in chapter.citations:
                    all_references.append({
                        "chapter": chapter.number,
                        "filename": cit.filename,
                        "quote": cit.quote,
                    })
            
            lines.append("")
            lines.append("")
        
        # Add References section at the end of the book (not inline or per-chapter)
        if all_references:
            lines.append("=" * 60)
            lines.append("REFERENCES")
            lines.append("=" * 60)
            lines.append("")
            
            current_chapter = None
            for ref in all_references:
                if ref["chapter"] != current_chapter:
                    if current_chapter is not None:
                        lines.append("")
                    lines.append(f"Chapter {ref['chapter']}:")
                    current_chapter = ref["chapter"]
                
                # Show source file and the quote (if not too long)
                quote_preview = ref["quote"][:100] + "..." if len(ref["quote"]) > 100 else ref["quote"]
                lines.append(f"  - {ref['filename']}: \"{quote_preview}\"")
            
            lines.append("")
        
        return '\n'.join(lines)
    
    # =========================================================================
    # HTML Export
    # =========================================================================
    
    def _export_html(self, chapters: list[Chapter], metadata: BookMetadata) -> str:
        """Export to HTML format."""
        html = f'''<!DOCTYPE html>
<html lang="{metadata.language}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <meta name="author" content="{metadata.author}">
    <meta name="description" content="{metadata.description}">
    <meta name="keywords" content="{', '.join(metadata.keywords)}">
    <title>{metadata.title}</title>
    <style>
        :root {{
            --primary-color: #1a1a2e;
            --secondary-color: #4a4a6a;
            --accent-color: #e94560;
            --bg-color: #fafafa;
            --text-color: #333;
        }}
        
        * {{
            box-sizing: border-box;
        }}
        
        body {{
            font-family: 'Georgia', 'Times New Roman', serif;
            line-height: 1.8;
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 20px;
            background: var(--bg-color);
            color: var(--text-color);
        }}
        
        .title-page {{
            text-align: center;
            padding: 100px 0;
            border-bottom: 2px solid var(--primary-color);
            margin-bottom: 60px;
        }}
        
        .title-page h1 {{
            font-size: 2.5em;
            color: var(--primary-color);
            margin-bottom: 20px;
        }}
        
        .title-page .author {{
            font-size: 1.3em;
            color: var(--secondary-color);
            margin-bottom: 30px;
        }}
        
        .title-page .meta {{
            font-size: 0.9em;
            color: #888;
        }}
        
        .toc {{
            margin: 40px 0;
            padding: 30px;
            background: #f0f0f0;
            border-radius: 8px;
        }}
        
        .toc h2 {{
            margin-top: 0;
            color: var(--primary-color);
        }}
        
        .toc ul {{
            list-style: none;
            padding: 0;
        }}
        
        .toc li {{
            padding: 8px 0;
            border-bottom: 1px solid #ddd;
        }}
        
        .toc a {{
            color: var(--secondary-color);
            text-decoration: none;
        }}
        
        .toc a:hover {{
            color: var(--accent-color);
        }}
        
        .chapter {{
            margin: 60px 0;
            padding-top: 40px;
            border-top: 1px solid #ddd;
        }}
        
        .chapter h1 {{
            color: var(--primary-color);
            font-size: 1.8em;
            margin-bottom: 30px;
        }}
        
        .chapter h2 {{
            color: var(--secondary-color);
            font-size: 1.3em;
            margin-top: 40px;
            margin-bottom: 20px;
        }}
        
        .chapter p {{
            text-align: justify;
            margin: 16px 0;
        }}
        
        .citation {{
            color: var(--secondary-color);
            font-size: 0.9em;
        }}
        
        strong {{
            color: var(--primary-color);
        }}
        
        footer {{
            text-align: center;
            margin-top: 80px;
            padding-top: 40px;
            border-top: 2px solid var(--primary-color);
            color: #888;
            font-size: 0.9em;
        }}
    </style>
</head>
<body>
    <div class="title-page">
        <h1>{metadata.title}</h1>
        <div class="author">by {metadata.author}</div>
        <div class="description">{metadata.description}</div>
        <div class="meta">{metadata.publisher} • {metadata.created_date}</div>
    </div>
    
    <nav class="toc">
        <h2>Table of Contents</h2>
        <ul>
'''
        
        # TOC entries
        for chapter in chapters:
            html += f'            <li><a href="#chapter-{chapter.number}">Chapter {chapter.number}: {chapter.title}</a></li>\n'
        
        html += '''        </ul>
    </nav>
'''
        
        # Chapters
        for chapter in chapters:
            chapter_body, marker_notes = self._extract_citations_as_footnotes(
                chapter.content, marker_style="none"
            )
            notes = []
            if chapter.citations:
                for cit in chapter.citations:
                    filename = (cit.filename or "Unknown").strip()
                    quote = (cit.quote or "").strip()
                    notes.append(f"{filename} - \"{quote}\"" if quote else filename)
            else:
                notes = marker_notes
            content_html = self._markdown_to_html(chapter_body)
            html += f'''
    <article class="chapter" id="chapter-{chapter.number}">
        <h1>Chapter {chapter.number}: {chapter.title}</h1>
        {content_html}
        {self._render_html_notes(notes)}
    </article>
'''
        
        html += f'''
    <footer>
        <p>Generated by GhostLine • {metadata.created_date}</p>
    </footer>
</body>
</html>
'''
        
        return html
    
    # =========================================================================
    # Markdown Export
    # =========================================================================
    
    def _export_markdown(self, chapters: list[Chapter], metadata: BookMetadata) -> str:
        """Export to Markdown format."""
        lines = []
        # We intentionally do NOT add inline footnote markers in markdown output.
        # Citations are rendered in per-chapter Notes sections.
        
        # Title
        lines.append(f"# {metadata.title}")
        lines.append(f"## by {metadata.author}")
        lines.append("")
        
        if metadata.description:
            lines.append(f"*{metadata.description}*")
            lines.append("")
        
        lines.append(f"**Publisher:** {metadata.publisher}")
        lines.append(f"**Date:** {metadata.created_date}")
        lines.append("")
        lines.append("---")
        lines.append("")
        
        # Table of Contents
        lines.append("## Table of Contents")
        lines.append("")
        for chapter in chapters:
            anchor = f"chapter-{chapter.number}-{chapter.title.lower().replace(' ', '-')}"
            lines.append(f"- [Chapter {chapter.number}: {chapter.title}](#{anchor})")
        lines.append("")
        lines.append("---")
        lines.append("")
        
        # Chapters
        for chapter in chapters:
            lines.append(f"# Chapter {chapter.number}: {chapter.title}")
            lines.append("")
            lines.append(f"*Word count: {chapter.word_count}*")
            lines.append("")
            chapter_body, marker_notes = self._extract_citations_as_footnotes(
                chapter.content or "", marker_style="none"
            )
            lines.append(chapter_body.strip())
            lines.append("")

            notes = []
            if chapter.citations:
                for cit in chapter.citations:
                    filename = (cit.filename or "Unknown").strip()
                    quote = (cit.quote or "").strip()
                    notes.append(f"{filename} - \"{quote}\"" if quote else filename)
            else:
                notes = marker_notes

            if notes:
                lines.append("## Notes")
                lines.append("")
                for i, note in enumerate(notes, 1):
                    lines.append(f"{i}. {note}")
                lines.append("")

            lines.append("---")
            lines.append("")
        
        return '\n'.join(lines)


# =============================================================================
# Convenience functions
# =============================================================================

def export_book_from_db(
    task_id: str,
    output_dir: str,
    formats: list[ExportFormat] = None,
) -> dict[ExportFormat, str]:
    """
    Export a book from database by task ID.
    
    Args:
        task_id: The generation task ID
        output_dir: Directory to save exports
        formats: List of formats to export (default: all)
        
    Returns:
        Dict mapping format to output file path
    """
    import json
    from sqlalchemy.orm import Session
    
    # Import here to avoid circular imports
    from app.db.base import SessionLocal
    from app.models.generation_task import GenerationTask
    from app.models.project import Project
    
    db = SessionLocal()
    try:
        task = db.query(GenerationTask).filter(GenerationTask.id == task_id).first()
        if not task:
            raise ValueError(f"Task not found: {task_id}")
        
        project = db.query(Project).filter(Project.id == task.project_id).first()
        
        # Extract chapters from output_data
        output_data = task.output_data or {}
        workflow_state = output_data.get("workflow_state", {})
        chapters_data = workflow_state.get("chapters", [])
        
        if not chapters_data:
            raise ValueError("No chapters found in task output")
        
        # Convert to Chapter objects
        chapters = [
            Chapter(
                number=ch.get("number", i + 1),
                title=ch.get("title", f"Chapter {i + 1}"),
                content=ch.get("content", ""),
                word_count=ch.get("word_count", 0),
            )
            for i, ch in enumerate(chapters_data)
        ]
        
        # Create metadata
        metadata = BookMetadata(
            title=project.title if project else "Untitled Book",
            author=project.author_name if project and hasattr(project, 'author_name') else "Anonymous",
            description=project.description if project else "",
        )
        
        # Export
        service = BookExportService(output_dir)
        
        if formats:
            results = {}
            for fmt in formats:
                try:
                    base = service._sanitize_filename(metadata.title)
                    path = Path(output_dir) / f"{base}.{fmt.value}"
                    service.export(chapters, metadata, fmt, str(path))
                    results[fmt] = str(path)
                except Exception as e:
                    results[fmt] = f"ERROR: {e}"
            return results
        else:
            return service.export_all(chapters, metadata, output_dir)
    
    finally:
        db.close()

